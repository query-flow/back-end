# app/main.py
from __future__ import annotations
import os, re, time, hashlib, base64
from uuid import uuid4
from typing import Any, Dict, List, Optional, Set, Tuple
from urllib.parse import urlencode, quote
from pathlib import Path
from io import BytesIO

# === Matplotlib headless (evita NSWindow em macOS/servers) ===
os.environ.setdefault("MPLBACKEND", "Agg")
try:
    import matplotlib
    matplotlib.use("Agg", force=True)
except Exception:
    pass

from fastapi import FastAPI, HTTPException, Depends, Header, UploadFile, File, Form
from pydantic import BaseModel, Field
from dotenv import load_dotenv
import httpx

# SQLAlchemy (config DB)
from sqlalchemy import (
    create_engine,
    text as sqltext,
    bindparam,
    String, Integer, Text, ForeignKey, JSON, UniqueConstraint
)
from sqlalchemy.orm import (
    DeclarativeBase, Mapped, mapped_column, relationship,
    sessionmaker, Session
)
from sqlalchemy.engine import make_url, Connection
from sqlalchemy.exc import IntegrityError, MultipleResultsFound

# Criptografia
from cryptography.fernet import Fernet, InvalidToken


# =========================================
# .env
# =========================================
DOTENV_PATH = (Path(__file__).resolve().parent.parent / ".env")
load_dotenv(dotenv_path=DOTENV_PATH, override=True)

AZURE_OPENAI_ENDPOINT    = os.getenv("AZURE_OPENAI_ENDPOINT", "").strip().rstrip("/")
AZURE_OPENAI_API_KEY     = os.getenv("AZURE_OPENAI_API_KEY", "").strip()
AZURE_OPENAI_DEPLOYMENT  = os.getenv("AZURE_OPENAI_DEPLOYMENT", "").strip()
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION", "2025-01-01-preview").strip()
DISABLE_AZURE_LLM        = os.getenv("DISABLE_AZURE_LLM", "0").strip() in {"1","true","True","YES","yes"}

CONFIG_DB_URL            = os.getenv("CONFIG_DB_URL", "mysql+pymysql://user:pass@127.0.0.1:3306/empresas?charset=utf8mb4").strip()
FERNET_KEY_B64           = os.getenv("FERNET_KEY", "").strip()
if not FERNET_KEY_B64:
    raise RuntimeError("FERNET_KEY ausente no .env. Gere com: from cryptography.fernet import Fernet; Fernet.generate_key().decode()")
fernet = Fernet(FERNET_KEY_B64.encode())

def encrypt_str(s: str) -> str:
    return fernet.encrypt(s.encode()).decode()

def decrypt_str(s: str) -> str:
    try:
        return fernet.decrypt(s.encode()).decode()
    except InvalidToken as e:
        raise HTTPException(status_code=400, detail="Não foi possível decifrar a senha armazenada.") from e


app = FastAPI(title="NL→SQL Multi-Org (MySQL) + RBAC + Bootstrap + Docs + Insights")


# =========================================
# ORM (banco de configuração)
# =========================================
class Base(DeclarativeBase):
    pass

class Org(Base):
    __tablename__ = "orgs"
    id: Mapped[str]     = mapped_column(String(36), primary_key=True)
    name: Mapped[str]   = mapped_column(String(120), unique=True)
    status: Mapped[str] = mapped_column(String(20), default="active")

    conn: Mapped["OrgDbConnection"] = relationship(back_populates="org", uselist=False, cascade="all, delete-orphan")
    schemas: Mapped[List["OrgAllowedSchema"]] = relationship(back_populates="org", cascade="all, delete-orphan")
    docs: Mapped[List["BizDocument"]] = relationship(back_populates="org", cascade="all, delete-orphan")
    members: Mapped[List["OrgMember"]] = relationship(back_populates="org", cascade="all, delete-orphan")

class OrgDbConnection(Base):
    __tablename__ = "org_db_connections"
    org_id: Mapped[str]        = mapped_column(String(36), ForeignKey("orgs.id"), primary_key=True)
    driver: Mapped[str]        = mapped_column(String(40))
    host: Mapped[str]          = mapped_column(String(255))
    port: Mapped[int]          = mapped_column(Integer, default=3306)
    username: Mapped[str]      = mapped_column(String(255))
    password_enc: Mapped[str]  = mapped_column(Text)
    database_name: Mapped[str] = mapped_column(String(255))
    options_json: Mapped[dict] = mapped_column(JSON, default={})

    org: Mapped[Org] = relationship(back_populates="conn")

class OrgAllowedSchema(Base):
    __tablename__ = "org_allowed_schemas"
    org_id: Mapped[str]      = mapped_column(String(36), ForeignKey("orgs.id"), primary_key=True)
    schema_name: Mapped[str] = mapped_column(String(255), primary_key=True)
    org: Mapped[Org]         = relationship(back_populates="schemas")

class BizDocument(Base):
    __tablename__ = "biz_documents"
    id: Mapped[int]            = mapped_column(Integer, primary_key=True, autoincrement=True)
    org_id: Mapped[str]        = mapped_column(String(36), ForeignKey("orgs.id"))
    title: Mapped[str]         = mapped_column(String(255))
    # Não armazenamos arquivo nem URL: só metadados ricos extraídos
    metadata_json: Mapped[dict]= mapped_column(JSON, default={})
    org: Mapped[Org]           = relationship(back_populates="docs")

class QueryAudit(Base):
    __tablename__ = "query_audit"
    id: Mapped[int]            = mapped_column(Integer, primary_key=True, autoincrement=True)
    org_id: Mapped[str]        = mapped_column(String(36))
    schema_used: Mapped[str]   = mapped_column(String(255))
    prompt_snip: Mapped[str]   = mapped_column(String(500))
    sql_text: Mapped[str]      = mapped_column(Text)
    row_count: Mapped[Optional[int]] = mapped_column(Integer, nullable=True)
    duration_ms: Mapped[Optional[int]] = mapped_column(Integer, nullable=True)

# === RBAC ===
class User(Base):
    __tablename__ = "users"
    __table_args__ = (UniqueConstraint("api_key_sha", name="uq_users_api_key_sha"),)
    id: Mapped[str]        = mapped_column(String(36), primary_key=True)
    name: Mapped[str]      = mapped_column(String(120))
    email: Mapped[str]     = mapped_column(String(255), unique=True)
    role: Mapped[str]      = mapped_column(String(10))  # 'admin' | 'user'
    api_key_sha: Mapped[str] = mapped_column(String(64))  # sha256 da API key (UNIQUE via constraint)
    org_links: Mapped[List["OrgMember"]] = relationship(back_populates="user", cascade="all, delete-orphan")

class OrgMember(Base):
    __tablename__ = "org_members"
    user_id: Mapped[str]   = mapped_column(String(36), ForeignKey("users.id"), primary_key=True)
    org_id:  Mapped[str]   = mapped_column(String(36), ForeignKey("orgs.id"),  primary_key=True)
    role_in_org: Mapped[str]= mapped_column(String(20), default="member")  # 'member'|'analyst'|'admin_org' etc.

    user: Mapped[User] = relationship(back_populates="org_links")
    org:  Mapped[Org]  = relationship(back_populates="members")


engine_cfg = create_engine(CONFIG_DB_URL, pool_pre_ping=True, future=True)
SessionCfg = sessionmaker(bind=engine_cfg, autoflush=False, autocommit=False, future=True)

def sha256_hex(s: str) -> str:
    return hashlib.sha256(s.encode()).hexdigest()

@app.on_event("startup")
def _startup():
    Base.metadata.create_all(engine_cfg)
    # Seed opcional de superadmin via .env
    sa_name  = os.getenv("SUPERADMIN_NAME", "").strip()
    sa_email = os.getenv("SUPERADMIN_EMAIL", "").strip()
    sa_key   = os.getenv("SUPERADMIN_API_KEY", "").strip()
    if sa_email and sa_key:
        with SessionCfg() as s:
            exists = s.query(User).filter_by(email=sa_email).one_or_none()
            if not exists:
                s.add(User(
                    id=uuid4().hex[:12],
                    name=sa_name or "Super Admin",
                    email=sa_email,
                    role="admin",
                    api_key_sha=sha256_hex(sa_key)
                ))
                s.commit()


# =========================================
# Utilitários (URL de DB do cliente)
# =========================================
def parse_database_url(url_str: str) -> dict:
    try:
        u = make_url(url_str)
    except Exception:
        raise HTTPException(status_code=400, detail="database_url inválida.")
    if not u.database:
        raise HTTPException(status_code=400, detail="A database_url deve incluir um DB/schema (ex.: ...:3306/sakila?charset=utf8mb4).")
    return {
        "driver": u.drivername,
        "host": u.host or "127.0.0.1",
        "port": int(u.port or 3306),
        "username": u.username or "",
        "password": u.password or "",
        "database_name": u.database,
        "options": dict(u.query or {}),
    }

def build_sqlalchemy_url(driver: str, host: str, port: int, username: str,
                         password_plain: str, database: str, options: Optional[dict]) -> str:
    pwd_enc = quote(password_plain, safe="")
    qs = "?" + urlencode(options or {}) if options else ""
    return f"{driver}://{username}:{pwd_enc}@{host}:{port}/{database}{qs}"


# =========================================
# NL→SQL (Azure OpenAI) + fallback
# =========================================
SYSTEM_PROMPT_TEMPLATE = """Você é um tradutor NL→SQL no dialeto {dialeto}. Regras obrigatórias:
- Gere SOMENTE um SELECT SQL válido (sem comentários, sem ```).
- Use apenas tabelas/colunas do esquema fornecido.
- Prefira JOINs com PK/FK explícitas.
- NUNCA modifique dados (sem INSERT/UPDATE/DELETE/DDL).
- Se o usuário não pedir limite explícito, inclua LIMIT {limit}.
- Dialeto alvo: {dialeto}.
"""

PROMPT_BASE = """{esquema}

Pergunta do usuário:
{pergunta}

Responda SOMENTE com o SQL válido (sem explicações)."""

def _azure_chat_url() -> str:
    return f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{AZURE_OPENAI_DEPLOYMENT}/chat/completions?api-version={AZURE_OPENAI_API_VERSION}"

def chamar_llm_azure(prompt_usuario: str, limit: int = 100, dialeto: str = "MySQL") -> str:
    if DISABLE_AZURE_LLM or not AZURE_OPENAI_API_KEY or not AZURE_OPENAI_ENDPOINT or not AZURE_OPENAI_DEPLOYMENT:
        return f"SELECT 1 AS ok LIMIT {limit};"
    system_prompt = SYSTEM_PROMPT_TEMPLATE.format(dialeto=dialeto, limit=limit)
    payload = {"messages":[{"role":"system","content":system_prompt},{"role":"user","content":prompt_usuario}],
               "temperature":0.1, "max_tokens":800, "top_p":0.95}
    headers = {"Content-Type":"application/json","api-key":AZURE_OPENAI_API_KEY}
    url = _azure_chat_url()
    try:
        with httpx.Client(timeout=httpx.Timeout(30.0, connect=10.0)) as client:
            r = client.post(url, headers=headers, json=payload)
            r.raise_for_status()
            data = r.json()
        content = data["choices"][0]["message"]["content"].strip()
        if content.startswith("```"):
            content = content.strip("`")
            if content.lower().startswith("sql"):
                content = content[3:].lstrip()
        return content.strip()
    except Exception:
        return f"SELECT 1 AS ok LIMIT {limit};"


# =========================================
# Reflection / Catálogo (DB atual)
# =========================================
def _catalog_for_current_db(conn: Connection, db_name: str) -> Dict[str, Any]:
    tables: Dict[str, Any] = {}

    cols = conn.execute(sqltext("""
        SELECT TABLE_NAME, COLUMN_NAME, DATA_TYPE, COLUMN_KEY
        FROM information_schema.COLUMNS
        WHERE TABLE_SCHEMA = :db
        ORDER BY TABLE_NAME, ORDINAL_POSITION
    """), {"db": db_name}).mappings()
    for row in cols:
        t = row["TABLE_NAME"]
        tables.setdefault(t, {"columns": [], "pks": set(), "fks": []})
        tables[t]["columns"].append({"name": row["COLUMN_NAME"], "type": row["DATA_TYPE"]})
        if row["COLUMN_KEY"] == "PRI":
            tables[t]["pks"].add(row["COLUMN_NAME"])

    fks = conn.execute(sqltext("""
        SELECT TABLE_NAME, COLUMN_NAME, REFERENCED_TABLE_NAME, REFERENCED_COLUMN_NAME
        FROM information_schema.KEY_COLUMN_USAGE
        WHERE TABLE_SCHEMA = :db AND REFERENCED_TABLE_NAME IS NOT NULL
        ORDER BY TABLE_NAME, COLUMN_NAME
    """), {"db": db_name}).mappings()
    for row in fks:
        t = row["TABLE_NAME"]
        tables.setdefault(t, {"columns": [], "pks": set(), "fks": []})
        tables[t]["fks"].append({
            "col": row["COLUMN_NAME"],
            "ref_table": row["REFERENCED_TABLE_NAME"],
            "ref_col": row["REFERENCED_COLUMN_NAME"]
        })

    return {"db": db_name, "tables": tables}

def _esquema_resumido(catalog: Dict[str, Any], max_chars: int = 4000) -> str:
    linhas: List[str] = []
    for t, meta in catalog["tables"].items():
        cols = ", ".join([f'{c["name"]}:{c["type"]}' for c in meta["columns"]][:24])
        linhas.append(f"- {t}({cols})")
    texto = "Esquema disponível:\n" + "\n".join(linhas)
    return texto[:max_chars]


# =========================================
# Guardrails
# =========================================
SQL_PERIGOSO = re.compile(r"\b(INSERT|UPDATE|DELETE|MERGE|ALTER|DROP|CREATE|TRUNCATE|GRANT|REVOKE)\b", re.I)

def proteger_sql_singledb(sql: str, catalog: Dict[str, Any], db_name: str, max_linhas: int) -> str:
    if SQL_PERIGOSO.search(sql):
        raise HTTPException(status_code=400, detail="SQL com comandos de escrita/DDL não é permitido.")

    refs = re.findall(r'(?:from|join)\s+((?:[`"]?[a-zA-Z0-9_]+[`"]?\.)?[`"]?[a-zA-Z0-9_]+[`"]?)', sql, flags=re.I)

    def split_ref(r: str) -> Tuple[Optional[str], str]:
        r = r.strip('`"')
        parts = r.split(".")
        if len(parts) == 2:
            return parts[0].lower(), parts[1].lower()
        return None, parts[0].lower()

    other_dbs: Set[str] = set()
    tables_used: Set[str] = set()
    for ref in refs:
        db, tb = split_ref(ref)
        if db and db != db_name.lower():
            other_dbs.add(db)
        tables_used.add(tb)

    if other_dbs:
        raise HTTPException(status_code=400, detail=f"Tabelas desconhecidas (multi-DB não permitido): {other_dbs}")

    known = {k.lower() for k in catalog["tables"].keys()}
    unknown = {t for t in tables_used if t not in known}
    if unknown:
        raise HTTPException(status_code=400, detail=f"Tabela(s) não encontrada(s) em {db_name}: {unknown}")

    if re.search(r"\blimit\b", sql, flags=re.I) is None:
        sql += f"\nLIMIT {max_linhas}"
    if not sql.strip().endswith(";"):
        sql += ";"
    return sql


# =========================================
# Execução
# =========================================
def executar_sql_readonly_on_conn(conn: Connection, sql: str) -> Dict[str, Any]:
    rs = conn.execute(sqltext(sql))
    cols = list(rs.keys())
    dados = [dict(zip(cols, row)) for row in rs]
    return {"colunas": cols, "dados": dados}


# =========================================
# Schemas (requests)
# =========================================
class AdminOrgCreate(BaseModel):
    name: str = Field(..., examples=["Empresa X"])
    database_url: str = Field(..., description="SQLAlchemy URL com DB (schema) obrigatório")
    allowed_schemas: List[str] = Field(..., min_items=1)
    documents: List[Dict[str, Any]] = Field(default_factory=list)

class AdminOrgResponse(BaseModel):
    org_id: str
    name: str
    allowed_schemas: List[str]

class AdminUserCreate(BaseModel):
    name: str
    email: str
    role: str = Field(..., pattern="^(admin|user)$")
    api_key_plain: str = Field(..., min_length=16)
    # vinculação automática opcional
    org_id: Optional[str] = None
    org_role: Optional[str] = Field(default="member")

class AdminUserResponse(BaseModel):
    user_id: str
    name: str
    email: str
    role: str

class AdminOrgMemberAdd(BaseModel):
    user_id: str
    role_in_org: str = "member"

class PerguntaOrg(BaseModel):
    org_id: str
    pergunta: str
    max_linhas: int = 100
    enrich: bool = True

class PerguntaDireta(BaseModel):
    database_url: str
    pergunta: str
    max_linhas: int = 100

class AdminDocManualCreate(BaseModel):
    title: str
    metadata_json: Dict[str, Any]


# =========================================
# Auth dependencies
# =========================================
class AuthedUser(BaseModel):
    id: str
    email: str
    role: str

def get_db():
    with SessionCfg() as s:
        yield s

def auth_required(x_api_key: str | None = Header(default=None, alias="X-API-Key"),
                  db: Session = Depends(get_db)) -> AuthedUser:
    if not x_api_key:
        raise HTTPException(status_code=401, detail="X-API-Key ausente.")
    try:
        user = db.query(User).filter_by(api_key_sha=sha256_hex(x_api_key)).one_or_none()
    except MultipleResultsFound:
        raise HTTPException(status_code=401, detail="API Key ambígua (duplicada). Rotacione a chave ou contate o admin.")
    if not user:
        raise HTTPException(status_code=401, detail="API Key inválida.")
    return AuthedUser(id=user.id, email=user.email, role=user.role)

def require_admin(user: AuthedUser = Depends(auth_required)) -> AuthedUser:
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Acesso restrito a admin.")
    return user

def require_user_or_admin(user: AuthedUser = Depends(auth_required)) -> AuthedUser:
    if user.role not in ("admin", "user"):
        raise HTTPException(status_code=403, detail="Acesso negado.")
    return user

def require_org_access(org_id: str, user: AuthedUser, db: Session) -> None:
    if user.role == "admin":
        return
    link = db.query(OrgMember).filter_by(user_id=user.id, org_id=org_id).first()
    if not link:
        raise HTTPException(status_code=403, detail="Sem acesso a esta organização.")


# =========================================
# Endpoints ADMIN (protegidos)
# =========================================
@app.post("/admin/orgs", response_model=AdminOrgResponse)
def admin_create_org(payload: AdminOrgCreate, _u: AuthedUser = Depends(require_admin)):
    try:
        parts = parse_database_url(payload.database_url)
        if not parts["username"] or not parts["password"]:
            raise HTTPException(status_code=400, detail="database_url deve conter usuário e senha.")

        org_id = uuid4().hex[:12]
        with SessionCfg() as db:
            if db.query(Org).filter(Org.name == payload.name).first():
                raise HTTPException(status_code=400, detail="Já existe uma organização com esse nome.")
            org = Org(id=org_id, name=payload.name, status="active")
            db.add(org)
            db.add(OrgDbConnection(
                org_id=org_id,
                driver=parts["driver"], host=parts["host"], port=parts["port"],
                username=parts["username"], password_enc=encrypt_str(parts["password"]),
                database_name=parts["database_name"], options_json=parts["options"]
            ))
            for s in payload.allowed_schemas:
                db.add(OrgAllowedSchema(org_id=org_id, schema_name=s))
            for d in payload.documents:
                db.add(BizDocument(org_id=org_id, title=d["title"], metadata_json=d.get("metadata_json", {})))
            db.commit()
        return AdminOrgResponse(org_id=org_id, name=payload.name, allowed_schemas=payload.allowed_schemas)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/admin/orgs/{org_id}", response_model=AdminOrgResponse)
def admin_get_org(org_id: str, _u: AuthedUser = Depends(require_admin)):
    with SessionCfg() as db:
        org = db.get(Org, org_id)
        if not org:
            raise HTTPException(status_code=404, detail="org_id não encontrado.")
        schemas = [s.schema_name for s in org.schemas]
        return AdminOrgResponse(org_id=org.id, name=org.name, allowed_schemas=schemas)

@app.post("/admin/orgs/{org_id}/test-connection")
def admin_test_connection(org_id: str, _u: AuthedUser = Depends(require_admin)):
    with SessionCfg() as db:
        org = db.get(Org, org_id)
        if not (org and org.conn):
            raise HTTPException(status_code=404, detail="org_id não encontrado.")
        pwd = decrypt_str(org.conn.password_enc)
        db_url = build_sqlalchemy_url(
            org.conn.driver, org.conn.host, org.conn.port,
            org.conn.username, pwd, org.conn.database_name, org.conn.options_json
        )
    eng = create_engine(db_url, pool_pre_ping=True, future=True)
    try:
        with eng.connect() as c:
            cur = c.execute(sqltext("SELECT DATABASE()")).scalar()
            one = c.execute(sqltext("SELECT 1")).scalar()
        return {"ok": True, "database_corrente": cur, "select_1": one}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

# --- Users (admin) ---
@app.post("/admin/users", response_model=AdminUserResponse)
def admin_create_user(payload: AdminUserCreate, _u: AuthedUser = Depends(require_admin)):
    with SessionCfg() as db:
        if db.query(User).filter_by(email=payload.email).one_or_none():
            raise HTTPException(status_code=400, detail="Email já cadastrado.")
        if db.query(User).filter_by(api_key_sha=sha256_hex(payload.api_key_plain)).one_or_none():
            raise HTTPException(status_code=400, detail="API key já está em uso. Gere uma diferente.")
        user = User(
            id=uuid4().hex[:12],
            name=payload.name,
            email=payload.email,
            role=payload.role,
            api_key_sha=sha256_hex(payload.api_key_plain),
        )
        db.add(user)
        # vinculação automática opcional
        if payload.org_id:
            org = db.get(Org, payload.org_id)
            if not org:
                raise HTTPException(status_code=404, detail="org_id para vincular não encontrada.")
            db.add(OrgMember(user_id=user.id, org_id=payload.org_id, role_in_org=payload.org_role or "member"))
        db.commit()
        return AdminUserResponse(user_id=user.id, name=user.name, email=user.email, role=user.role)

@app.post("/admin/orgs/{org_id}/members")
def admin_add_member(org_id: str, payload: AdminOrgMemberAdd, _u: AuthedUser = Depends(require_admin)):
    with SessionCfg() as db:
        org = db.get(Org, org_id)
        user = db.get(User, payload.user_id)
        if not org:
            raise HTTPException(status_code=404, detail="org não encontrada.")
        if not user:
            raise HTTPException(status_code=404, detail="user não encontrado.")
        link = db.query(OrgMember).filter_by(org_id=org_id, user_id=user.id).one_or_none()
        if link:
            link.role_in_org = payload.role_in_org
        else:
            db.add(OrgMember(user_id=user.id, org_id=org_id, role_in_org=payload.role_in_org))
        db.commit()
        return {"ok": True, "org_id": org_id, "user_id": user.id, "role_in_org": payload.role_in_org}

@app.delete("/admin/users/{user_id}")
def admin_delete_user(user_id: str, _u: AuthedUser = Depends(require_admin)):
    with SessionCfg() as db:
        user = db.get(User, user_id)
        if not user:
            raise HTTPException(status_code=404, detail="user não encontrado.")
        db.delete(user)
        db.commit()
        return {"ok": True, "deleted_user_id": user_id}


# =========================================
# Roteamento automático de schema (índice + LLM classificador)
# =========================================
_SCHEMA_INDEX_CACHE: dict[str, dict[str, set[str]]] = {}
_SCHEMA_INDEX_TTL: dict[str, float] = {}
_SCHEMA_INDEX_MAX_AGE = 300.0  # 5 min

def _normalize_tokens(*parts: str) -> set[str]:
    toks: set[str] = set()
    for p in parts:
        for t in re.findall(r"[a-zA-Z0-9_]+", (p or "").lower()):
            toks.add(t)
    return toks

def build_schema_index(conn: Connection, allowed_schemas: list[str]) -> dict[str, set[str]]:
    rows = conn.execute(
        sqltext("""
            SELECT TABLE_SCHEMA, TABLE_NAME, COLUMN_NAME
            FROM information_schema.COLUMNS
            WHERE TABLE_SCHEMA IN :schemas
        """).bindparams(bindparam("schemas", expanding=True)),
        {"schemas": allowed_schemas}
    ).fetchall()

    index: dict[str, set[str]] = {s: set() for s in allowed_schemas}
    for schema, table, col in rows:
        index[schema].add((table or "").lower())
        index[schema].add((col or "").lower())
    return index

def get_schema_index_for_org(org_id: str, base_db_url_with_default: str, allowed: list[str]) -> dict[str, set[str]]:
    now = time.time()
    if (org_id in _SCHEMA_INDEX_CACHE) and (now - _SCHEMA_INDEX_TTL.get(org_id, 0) < _SCHEMA_INDEX_MAX_AGE):
        return _SCHEMA_INDEX_CACHE[org_id]
    eng = create_engine(base_db_url_with_default, pool_pre_ping=True, future=True)
    with eng.connect() as conn:
        idx = build_schema_index(conn, allowed)
    _SCHEMA_INDEX_CACHE[org_id] = idx
    _SCHEMA_INDEX_TTL[org_id] = now
    return idx

def rank_schemas_by_overlap(schema_index: dict[str, set[str]], pergunta: str) -> list[tuple[str, int]]:
    q_toks = _normalize_tokens(pergunta)
    scored: list[tuple[str, int]] = []
    for schema, tokens in schema_index.items():
        score = len(q_toks & tokens)
        scored.append((schema, score))
    scored.sort(key=lambda x: (-x[1], x[0].lower()))
    return scored

def ask_llm_pick_schema(allowed: list[str], pergunta: str) -> Optional[str]:
    if DISABLE_AZURE_LLM or not AZURE_OPENAI_API_KEY or not AZURE_OPENAI_ENDPOINT or not AZURE_OPENAI_DEPLOYMENT:
        return None
    system = "Você escolhe UM schema dentre os permitidos. Responda com uma única palavra (nome exato), sem explicações."
    user = f"Schemas permitidos: {', '.join(allowed)}\nPergunta: {pergunta}\nResponda apenas com o schema."
    payload = {"messages":[{"role":"system","content":system},{"role":"user","content":user}],
               "temperature":0.0, "max_tokens":10}
    headers = {"Content-Type":"application/json","api-key":AZURE_OPENAI_API_KEY}
    url = _azure_chat_url()
    try:
        with httpx.Client(timeout=httpx.Timeout(10.0, connect=5.0)) as cli:
            r = cli.post(url, headers=headers, json=payload)
            r.raise_for_status()
            data = r.json()
        choice = data["choices"][0]["message"]["content"].strip()
        cand = re.findall(r"[a-zA-Z0-9_]+", choice)
        if not cand: return None
        name = cand[0].lower()
        for s in allowed:
            if s.lower() == name:
                return s
        return None
    except Exception:
        return None


# =========================================
# Insights helpers
# =========================================
def _collect_biz_context_for_org(org: Org) -> str:
    if not org.docs:
        return "Sem documentos de negócio cadastrados."
    partes = []
    for d in org.docs:
        md = d.metadata_json or {}
        md_txt = "; ".join(f"{k}: {v}" for k, v in md.items())
        partes.append(f"- {d.title} ({md_txt})" if md_txt else f"- {d.title}")
    return "Documentos de negócio cadastrados:\n" + "\n".join(partes)

def _pick_chart_axes(resultado: Dict[str, Any]) -> Optional[Tuple[List[str], List[float], str]]:
    cols = resultado.get("colunas", [])
    rows = resultado.get("dados", [])
    if not cols or not rows:
        return None

    # Tenta detectar primeira coluna categórica (string-ish)
    cat_idx = None
    for i, c in enumerate(cols):
        # se a maioria dos valores é string/não numérico, assume categórica
        sample = [r.get(c) for r in rows[:10]]
        non_num = 0
        for v in sample:
            try:
                float(v)
                # numérico
            except:
                non_num += 1
        if non_num >= max(1, len(sample)//2):
            cat_idx = i
            break

    # Se não achou categórica, usa a coluna 0 como rótulo
    if cat_idx is None:
        cat_idx = 0

    # Tenta achar uma coluna numérica diferente da categórica
    num_idx = None
    for j, c in enumerate(cols):
        if j == cat_idx: continue
        ok = True
        for v in [r.get(c) for r in rows[:10]]:
            try:
                float(v)
            except:
                ok = False
                break
        if ok:
            num_idx = j
            break

    if num_idx is None:
        return None

    labels = [str(r.get(cols[cat_idx])) for r in rows]
    values = []
    for r in rows:
        try:
            values.append(float(r.get(cols[num_idx]) or 0))
        except:
            values.append(0.0)

    title = f"{cols[num_idx]} por {cols[cat_idx]}"
    return labels, values, title

def _make_bar_chart_base64_generic(resultado: Dict[str, Any]) -> Optional[str]:
    try:
        import matplotlib.pyplot as plt
    except Exception:
        return None

    picked = _pick_chart_axes(resultado)
    if not picked:
        return None
    labels, values, title = picked

    fig = plt.figure()
    plt.bar(labels, values)
    plt.xticks(rotation=45, ha="right")
    plt.title(title)
    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    return base64.b64encode(buf.getvalue()).decode()

def _insights_from_llm(pergunta: str, resultado: Dict[str, Any], biz_context: str) -> str:
    preview_rows = resultado.get("dados", [])[:10]
    preview_cols = resultado.get("colunas", [])
    preview_text = f"Colunas: {preview_cols}\nAmostra (até 10 linhas): {preview_rows}"
    system = (
        "Você é um analista de BI. Explique resultados em linguagem de negócio, "
        "conectando com o contexto fornecido, sem jargões técnicos desnecessários. "
        "Seja objetivo."
    )
    user = (
        f"{biz_context}\n\n"
        f"Pergunta do usuário: {pergunta}\n\n"
        f"Prévia dos dados (tabela):\n{preview_text}\n\n"
        "Escreva uma síntese de até 8 linhas com o que é mais relevante, seguida de 3 pontos de atenção / próximos passos."
    )
    if DISABLE_AZURE_LLM:
        return "Insights desativados em ambiente de teste (DISABLE_AZURE_LLM=1)."
    payload = {
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user}
        ],
        "temperature": 0.2,
        "max_tokens": 500,
        "top_p": 0.9
    }
    headers = {"Content-Type": "application/json", "api-key": AZURE_OPENAI_API_KEY}
    url = _azure_chat_url()
    try:
        with httpx.Client(timeout=httpx.Timeout(30.0, connect=10.0)) as client:
            r = client.post(url, headers=headers, json=payload)
            r.raise_for_status()
            data = r.json()
        return data["choices"][0]["message"]["content"].strip()
    except Exception as e:
        return f"(Não foi possível gerar insights agora: {e})"


# =========================================
# Endpoint principal (usuário final)
# =========================================
@app.post("/perguntar_org")
def perguntar_org(p: PerguntaOrg,
                  u: AuthedUser = Depends(require_user_or_admin),
                  db: Session = Depends(get_db)):
    try:
        # autorização por organização
        require_org_access(p.org_id, u, db)

        # 1) Carrega config da org + contexto de negócio
        with SessionCfg() as s:
            org = s.get(Org, p.org_id)
            if not (org and org.conn):
                raise HTTPException(status_code=404, detail="org_id inválido.")
            allowed = [x.schema_name for x in org.schemas]
            if not allowed:
                raise HTTPException(status_code=400, detail="Org sem schemas permitidos.")
            biz_context = _collect_biz_context_for_org(org)

            pwd = decrypt_str(org.conn.password_enc)
            base_parts = (org.conn.driver, org.conn.host, org.conn.port, org.conn.username, pwd, org.conn.options_json, org.conn.database_name)

        # 2) Router automático de schema
        base_url_default_db = build_sqlalchemy_url(
            base_parts[0], base_parts[1], base_parts[2], base_parts[3], base_parts[4],
            base_parts[6], base_parts[5]
        )
        schema_index = get_schema_index_for_org(p.org_id, base_url_default_db, allowed)
        ranked = rank_schemas_by_overlap(schema_index, p.pergunta)
        best_by_overlap, top_score = ranked[0]
        top_ties = [s for s, sc in ranked if sc == top_score]
        if top_score == 0 or len(top_ties) > 1:
            picked = ask_llm_pick_schema(allowed, p.pergunta)
            preferred = picked or best_by_overlap
        else:
            preferred = best_by_overlap
        schema_try_order = [preferred] + [s for s in allowed if s != preferred]

        last_err: Optional[str] = None

        # 3) Tenta executar na ordem decidida
        for schema in schema_try_order:
            db_url = build_sqlalchemy_url(base_parts[0], base_parts[1], base_parts[2],
                                          base_parts[3], base_parts[4], schema, base_parts[5])
            eng = create_engine(db_url, pool_pre_ping=True, future=True)
            try:
                with eng.connect() as conn:
                    catalog = _catalog_for_current_db(conn, db_name=schema)
                    esquema_txt = _esquema_resumido(catalog)
                    prompt = PROMPT_BASE.format(esquema=esquema_txt, pergunta=p.pergunta)
                    sql = chamar_llm_azure(prompt_usuario=prompt, limit=p.max_linhas, dialeto="MySQL")

                    try:
                        sql_seguro = proteger_sql_singledb(sql, catalog, db_name=schema, max_linhas=p.max_linhas)
                    except HTTPException as e:
                        msg = str(e.detail)
                        if "Tabela(s) não encontrada(s)" in msg or "multi-DB" in msg:
                            last_err = f"[{schema}] {msg}"
                            continue
                        raise

                    t0 = time.time()
                    try:
                        resultado = executar_sql_readonly_on_conn(conn, sql_seguro)
                    except Exception as err:
                        corre = f"Esquema:\n{esquema_txt}\n\nErro:\n{err}\n\nCorrija (somente SELECT, LIMIT {p.max_linhas} se faltar):\n{sql_seguro}"
                        sql2 = chamar_llm_azure(corre, limit=p.max_linhas, dialeto="MySQL")
                        sql_seguro = proteger_sql_singledb(sql2, catalog, db_name=schema, max_linhas=p.max_linhas)
                        resultado = executar_sql_readonly_on_conn(conn, sql_seguro)
                    dur_ms = int((time.time() - t0) * 1000)

                # auditoria best-effort
                try:
                    with SessionCfg() as s:
                        s.add(QueryAudit(
                            org_id=p.org_id, schema_used=schema,
                            prompt_snip=p.pergunta[:500], sql_text=sql_seguro,
                            row_count=len(resultado["dados"]) if resultado and "dados" in resultado else None,
                            duration_ms=dur_ms
                        ))
                        s.commit()
                except Exception:
                    pass

                insights_payload = None
                if p.enrich:
                    summary = _insights_from_llm(p.pergunta, resultado, biz_context)
                    chart_b64 = _make_bar_chart_base64_generic(resultado)
                    chart = {"mime": "image/png", "base64": chart_b64} if chart_b64 else None
                    insights_payload = {"summary": summary, "chart": chart}

                return {
                    "org_id": p.org_id,
                    "schema_usado": schema,
                    "sql": sql_seguro,
                    "resultado": resultado,
                    "insights": insights_payload
                }

            except HTTPException as e:
                last_err = f"[{schema}] {e.detail}"
                continue
            except Exception as e:
                last_err = f"[{schema}] {e}"
                continue

        raise HTTPException(status_code=400, detail=last_err or "Falha ao executar em todos os schemas permitidos.")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# =========================================
# BOOTSTRAP PÚBLICO (self-signup do admin + org)
# =========================================
class PublicBootstrapOrg(BaseModel):
    org_name: str
    database_url: str
    allowed_schemas: List[str]
    admin_name: str
    admin_email: str
    admin_api_key: str  # o próprio admin define a chave dele (min 16 chars)

class PublicBootstrapResponse(BaseModel):
    org_id: str
    admin_user_id: str
    admin_email: str

@app.post("/public/bootstrap_org", response_model=PublicBootstrapResponse)
def public_bootstrap_org(p: PublicBootstrapOrg):
    if len(p.admin_api_key) < 16:
        raise HTTPException(status_code=400, detail="admin_api_key deve ter pelo menos 16 caracteres.")
    if not p.allowed_schemas:
        raise HTTPException(status_code=400, detail="allowed_schemas não pode ser vazio.")

    parts = parse_database_url(p.database_url)
    if not parts["username"] or not parts["password"]:
        raise HTTPException(status_code=400, detail="database_url deve conter usuário e senha.")

    org_id = uuid4().hex[:12]
    admin_id = uuid4().hex[:12]

    try:
        with SessionCfg() as s:
            if s.query(Org).filter(Org.name == p.org_name).first():
                raise HTTPException(status_code=400, detail="Já existe uma organização com esse nome.")
            if s.query(User).filter(User.email == p.admin_email).first():
                raise HTTPException(status_code=400, detail="Já existe um usuário com esse e-mail.")
            if s.query(User).filter(User.api_key_sha == sha256_hex(p.admin_api_key)).first():
                raise HTTPException(status_code=400, detail="API key já está em uso. Informe outra.")

            org = Org(id=org_id, name=p.org_name, status="active")
            s.add(org)

            s.add(OrgDbConnection(
                org_id=org_id,
                driver=parts["driver"], host=parts["host"], port=parts["port"],
                username=parts["username"], password_enc=encrypt_str(parts["password"]),
                database_name=parts["database_name"], options_json=parts["options"]
            ))

            for sch in p.allowed_schemas:
                s.add(OrgAllowedSchema(org_id=org_id, schema_name=sch))

            admin_user = User(
                id=admin_id,
                name=p.admin_name,
                email=p.admin_email,
                role="admin",
                api_key_sha=sha256_hex(p.admin_api_key),
            )
            s.add(admin_user)
            s.add(OrgMember(user_id=admin_id, org_id=org_id, role_in_org="admin_org"))

            s.commit()

        return PublicBootstrapResponse(
            org_id=org_id, admin_user_id=admin_id, admin_email=p.admin_email
        )

    except IntegrityError as ie:
        raise HTTPException(status_code=400, detail=f"Violação de integridade no banco de configuração: {str(ie.orig)}") from ie
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Falha no bootstrap: {e}")


# =========================================
# Documentos (admin)
# =========================================
def _extract_text_from_upload(file: UploadFile) -> str:
    content = file.file.read()
    # tenta por extensão/mimetype
    name = (file.filename or "").lower()
    ctype = (file.content_type or "").lower()
    text = ""

    def _safe_decode(b: bytes) -> str:
        for enc in ("utf-8", "latin-1", "utf-16"):
            try:
                return b.decode(enc)
            except:
                continue
        return ""

    if name.endswith(".txt") or ctype.startswith("text/"):
        text = _safe_decode(content)
    elif name.endswith(".pdf") or "pdf" in ctype:
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(BytesIO(content))
            pages = []
            for p in reader.pages:
                pages.append(p.extract_text() or "")
            text = "\n".join(pages)
        except Exception:
            text = _safe_decode(content)
    elif name.endswith(".docx") or "officedocument.wordprocessingml.document" in ctype:
        try:
            import docx
            doc = docx.Document(BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            text = _safe_decode(content)
    else:
        # tentativa bruta
        text = _safe_decode(content)
    return (text or "").strip()

def _summarize_business_metadata(raw_text: str) -> Dict[str, Any]:
    """
    Gera metadados estruturados a partir do texto do documento.
    Se LLM estiver desabilitado, aplica heurísticas simples.
    """
    base_meta = {
        "summary": "",
        "kpis": [],
        "goals": [],
        "timeframe": "",
        "notes": "",
        "source_kind": "uploaded_document"
    }
    if not raw_text:
        base_meta["summary"] = "Documento vazio ou ilegível."
        return base_meta

    if DISABLE_AZURE_LLM or not AZURE_OPENAI_API_KEY or not AZURE_OPENAI_DEPLOYMENT or not AZURE_OPENAI_ENDPOINT:
        # heurística simples: pega primeiras linhas e tenta detectar números
        lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
        base_meta["summary"] = " ".join(lines[:5])[:600]
        nums = re.findall(r"\b\d+(?:[.,]\d+)?\b", raw_text)
        if nums:
            base_meta["kpis"] = [{"name": "valores_numericos_detectados", "values_sample": nums[:10]}]
        # metas heurística
        goals = []
        for m in re.findall(r"(meta|objetivo|target)[:\- ]+(.{0,120})", raw_text, flags=re.I):
            goals.append(m[1].strip())
        base_meta["goals"] = goals[:8]
        return base_meta

    # Com LLM
    system = (
        "Você extrai metadados de documentos de contexto de negócio, devolvendo JSON com as chaves: "
        "summary (string curta), kpis (lista de {name, formula?, current?, target?, unit?}), "
        "goals (lista de strings), timeframe (string), notes (string). Responda somente JSON."
    )
    user = f"Documento (texto puro):\n{raw_text[:12000]}"
    payload = {
        "messages":[{"role":"system","content":system},{"role":"user","content":user}],
        "temperature":0.1, "max_tokens":900, "top_p":0.95
    }
    headers = {"Content-Type":"application/json","api-key":AZURE_OPENAI_API_KEY}
    url = _azure_chat_url()
    try:
        with httpx.Client(timeout=httpx.Timeout(40.0, connect=10.0)) as client:
            r = client.post(url, headers=headers, json=payload)
            r.raise_for_status()
            data = r.json()
        content = data["choices"][0]["message"]["content"]
        # tenta parsear como JSON
        import json
        meta = json.loads(content)
        if isinstance(meta, dict):
            meta["source_kind"] = "uploaded_document"
            return meta
        base_meta["summary"] = str(meta)[:800]
        return base_meta
    except Exception as e:
        base_meta["summary"] = f"Falha ao usar LLM: {e}"
        return base_meta

@app.post("/admin/orgs/{org_id}/documents", response_model=dict)
def admin_add_document(org_id: str, payload: AdminDocManualCreate, _u: AuthedUser = Depends(require_admin)):
    with SessionCfg() as db:
        org = db.get(Org, org_id)
        if not org:
            raise HTTPException(status_code=404, detail="org não encontrada.")
        doc = BizDocument(org_id=org_id, title=payload.title, metadata_json=payload.metadata_json or {})
        db.add(doc)
        db.commit()
        return {"ok": True, "doc_id": doc.id, "org_id": org_id}

@app.post("/admin/orgs/{org_id}/documents/extract", response_model=dict)
def admin_extract_document(org_id: str,
                           title: str = Form(...),
                           file: UploadFile = File(...),
                           _u: AuthedUser = Depends(require_admin)):
    with SessionCfg() as db:
        org = db.get(Org, org_id)
        if not org:
            raise HTTPException(status_code=404, detail="org não encontrada.")
        raw_text = _extract_text_from_upload(file)
        meta = _summarize_business_metadata(raw_text)
        doc = BizDocument(org_id=org_id, title=title, metadata_json=meta)
        db.add(doc)
        db.commit()
        return {"ok": True, "doc_id": doc.id, "org_id": org_id, "title": title, "meta_preview": meta}

# =========================================
# Utilitários (protegidos para admin)
# =========================================
@app.post("/_debug_connect")
def debug_connect(p: PerguntaDireta, _u: AuthedUser = Depends(require_admin)):
    u = make_url(p.database_url)
    if not u.database:
        raise HTTPException(status_code=400, detail="Passe uma database_url com DB/schema.")
    eng = create_engine(p.database_url, pool_pre_ping=True, future=True)
    try:
        with eng.connect() as c:
            dbn = c.execute(sqltext("SELECT DATABASE()")).scalar()
            dbs = [r[0] for r in c.execute(sqltext("SHOW DATABASES"))]
            one = c.execute(sqltext("SELECT 1")).scalar()
        return {"ok": True, "database_corrente": dbn, "databases": dbs, "select_1": one}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/_env")
def _env(_u: AuthedUser = Depends(require_admin)):
    return {
        "loaded_dotenv_from": str(DOTENV_PATH),
        "DISABLE_AZURE_LLM": DISABLE_AZURE_LLM,
        "AZURE_OPENAI_ENDPOINT_set": bool(AZURE_OPENAI_ENDPOINT),
        "AZURE_OPENAI_DEPLOYMENT_set": bool(AZURE_OPENAI_DEPLOYMENT),
        "AZURE_OPENAI_API_KEY_set": bool(AZURE_OPENAI_API_KEY),
        "AZURE_OPENAI_API_VERSION": AZURE_OPENAI_API_VERSION,
    }
